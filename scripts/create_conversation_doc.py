# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

# Title
title = doc.add_heading('Claude Code Ярилцлага', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Date
date_para = doc.add_paragraph(f'Огноо: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1
doc.add_heading('1. IT Мониторинг ба Нууцлал', level=1)

doc.add_heading('Асуулт:', level=2)
doc.add_paragraph('Ажлын газрын IT нар ирээд өнөөдөр юу юу хийсэн, ямар ямар AI chat ашиглаж юу юу хийснийг мэдэх боломжтой юу?')

doc.add_heading('Хариулт:', level=2)
doc.add_paragraph('Ажлын газрын IT-ууд дараах зүйлсийг мэдэх боломжтой:')

table1 = doc.add_table(rows=7, cols=2)
table1.style = 'Table Grid'
headers = table1.rows[0].cells
headers[0].text = 'Төрөл'
headers[1].text = 'Тайлбар'

data1 = [
    ('Сүлжээний логууд', 'Ямар вэбсайт руу хандсан (claude.ai, chatgpt.com гэх мэт)'),
    ('DNS хүсэлтүүд', 'Бүх домайн хаягийн хүсэлтүүд'),
    ('Proxy/Firewall логууд', 'Интернэт траффик бүртгэл'),
    ('Endpoint мониторинг', 'Хэрвээ компьютерт monitoring software суулгасан бол'),
    ('Active Directory логууд', 'Нэвтрэлт, гарсан цаг'),
    ('Cloud backup', 'OneDrive, компанийн backup'),
]

for i, (col1, col2) in enumerate(data1, 1):
    row = table1.rows[i].cells
    row[0].text = col1
    row[1].text = col2

doc.add_paragraph()

# Section 2
doc.add_heading('2. Форматлахад устах уу?', level=1)

doc.add_heading('Асуулт:', level=2)
doc.add_paragraph('IT нарт өгөөд компьютерийг форматлуулчихвал яах вэ? Тэгсэний дараа шалгах боломжтой юу?')

doc.add_heading('Хариулт:', level=2)
doc.add_paragraph('Форматласан ч гэсэн IT-ууд зарим зүйлийг мэдэх боломжтой:')

table2 = doc.add_table(rows=8, cols=3)
table2.style = 'Table Grid'
headers2 = table2.rows[0].cells
headers2[0].text = 'Байршил'
headers2[1].text = 'Тайлбар'
headers2[2].text = 'Форматлахад устах уу?'

data2 = [
    ('Firewall/Proxy серверийн лог', 'Компанийн сервер дээр', 'Үгүй'),
    ('Active Directory лог', 'Domain controller дээр', 'Үгүй'),
    ('DNS сервер лог', 'Ямар сайт руу хандсан', 'Үгүй'),
    ('Network traffic лог', 'Сүлжээний төхөөрөмж дээр', 'Үгүй'),
    ('Email серверийн лог', 'Exchange/Gmail сервер', 'Үгүй'),
    ('Cloud backup', 'OneDrive, компанийн backup', 'Үгүй'),
    ('Локал файлууд, Browser history', 'Таны компьютер дотор', 'Тийм'),
]

for i, (col1, col2, col3) in enumerate(data2, 1):
    row = table2.rows[i].cells
    row[0].text = col1
    row[1].text = col2
    row[2].text = col3

doc.add_paragraph()
doc.add_paragraph('Дүгнэлт: Форматлах нь тусламжгүй - сүлжээний лог компанийн сервер дээр аль хэдийн бүртгэгдсэн байна.')

# Section 3
doc.add_heading('3. Зохиогчийн Эрхийн Асуудал', level=1)

doc.add_heading('Нөхцөл байдал:', level=2)
doc.add_paragraph('LIMS систем хөгжүүлсэн. Компаниас зохиогчийн эрх хамтран эзэмших ёстой гэж үздэг. Хөдөлмөрийн гэрээнд "ажиллах явцад ажилтаны санаачлан хийсэн бүтээл, хийсэн ажил компанийн өмч байна" гэсэн заалт байдаг.')

doc.add_heading('Ерөнхий байдал:', level=2)

table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Table Grid'
headers3 = table3.rows[0].cells
headers3[0].text = 'Нөхцөл'
headers3[1].text = 'Өмчлөгч'

data3 = [
    ('Ажлын цагаар + ажлын нөөцөөр хийсэн', 'Компани'),
    ('Ажлын үүргийн хүрээнд хийсэн', 'Компани'),
    ('Чөлөөт цагаар + хувийн нөөцөөр хийсэн', 'Ажилтан'),
]

for i, (col1, col2) in enumerate(data3, 1):
    row = table3.rows[i].cells
    row[0].text = col1
    row[1].text = col2

doc.add_paragraph()
doc.add_heading('Зөвлөгөө:', level=2)
doc.add_paragraph('1. Хөдөлмөрийн гэрээгээ нарийн уншаарай')
doc.add_paragraph('2. Хуульч/өмгөөлөгчтэй зөвлөлдөөрэй')
doc.add_paragraph('3. Нотлох баримт хадгалаарай')

# Section 4
doc.add_heading('4. Системээс Устгах Тухай', level=1)

doc.add_heading('Асуулт:', level=2)
doc.add_paragraph('Системээс устгах боломжтой юу?')

doc.add_heading('Хариулт:', level=2)
doc.add_paragraph('ҮГҮЙ, устгах хэрэггүй. Энэ нь маш муу санаа:')

table4 = doc.add_table(rows=6, cols=2)
table4.style = 'Table Grid'
headers4 = table4.rows[0].cells
headers4[0].text = 'Үр дагавар'
headers4[1].text = 'Тайлбар'

data4 = [
    ('Эрүүгийн хариуцлага', 'Компанийн өмч устгасан гэж үзэж болно'),
    ('Хөдөлмөрийн маргаан', 'Гэрээ зөрчсөн, нөхөн төлбөр нэхэмжилж болно'),
    ('Нэр хүнд', 'Ирээдүйн ажил олоход хүндрэл'),
    ('Backup-аас сэргээнэ', 'Устгасан ч IT-ууд сэргээж чадна'),
    ('Git history', 'Бүх өөрчлөлт хадгалагдсан байна'),
]

for i, (col1, col2) in enumerate(data4, 1):
    row = table4.rows[i].cells
    row[0].text = col1
    row[1].text = col2

doc.add_paragraph()
doc.add_paragraph('Устгах нь таны байр суурийг илүү муу болгоно. Тэвчээртэй байж, хуулийн зөв аргаар шийдээрэй.')

# Section 5
doc.add_heading('5. Лицензийн Бизнес Загвар', level=1)

doc.add_heading('Зорилго:', level=2)
doc.add_paragraph('LIMS программ амжилттай production-д нэвтэрсний дараа бусад нүүрсний лабораторид лиценз эзэмшүүлж мөнгө олох.')

doc.add_heading('Site License + Annual Fee загвар:', level=2)
doc.add_paragraph('Анхны суулгалт: ₮5,000,000 - ₮15,000,000')
doc.add_paragraph('Жилийн maintenance: ₮1,000,000 - ₮3,000,000')
doc.add_paragraph('Сургалт: ₮500,000 - ₮1,000,000')
doc.add_paragraph('Customization: Цагаар эсвэл тохиролцоно')

doc.add_heading('Лицензийн гэрээнд орох зүйлс:', level=2)
doc.add_paragraph('1. Тодорхойлолт - Програм хангамжийн нэр, хувилбар, лиценз олгогч, авагч')
doc.add_paragraph('2. Эрхийн хүрээ - Ашиглах эрх, хуулбарлах хориг')
doc.add_paragraph('3. Төлбөр - Анхны төлбөр, жилийн maintenance')
doc.add_paragraph('4. Хугацаа - Лицензийн хугацаа, сунгах нөхцөл')
doc.add_paragraph('5. Дэмжлэг, шинэчлэл - Техникийн дэмжлэг, response time')
doc.add_paragraph('6. Баталгаа, хязгаарлалт - Хариуцлагын хязгаарлалт')
doc.add_paragraph('7. Нууцлал - Өгөгдлийн нууцлал, эх кодын нууцлал')

# Section 6
doc.add_heading('6. Програм Хамгаалах Аргууд', level=1)

doc.add_heading('Асуудал:', level=2)
doc.add_paragraph('Том компаниуд өөрсдийн сервер дээр байршуулахыг хүсдэг. Тэдний IT нар программыг хуулж авах эрсдэлтэй.')

doc.add_heading('Шийдлүүд:', level=2)

table5 = doc.add_table(rows=5, cols=2)
table5.style = 'Table Grid'
headers5 = table5.rows[0].cells
headers5[0].text = 'Арга'
headers5[1].text = 'Тайлбар'

data5 = [
    ('Гэрээн дээр тулгуурлах', 'Том компаниуд нэр хүндээ хамгаална, гэрээ зөрчих эрсдэл авахгүй'),
    ('Жил бүр Maintenance зарах', 'Шинэ хувилбар, update, техникийн тусламж, сургалт'),
    ('Hardware донгл', 'USB түлхүүргүй бол ажиллахгүй'),
    ('Мэргэжлийн хамаарал', 'Зөвхөн та системийг мэддэг, засвар хийж чаддаг'),
]

for i, (col1, col2) in enumerate(data5, 1):
    row = table5.rows[i].cells
    row[0].text = col1
    row[1].text = col2

# Section 7
doc.add_heading('7. Хамгийн Тохиромжтой Стратеги', level=1)

doc.add_heading('Одоогийн байдал:', level=2)
doc.add_paragraph('- Ажлын газар: Energy Resources')
doc.add_paragraph('- Зорилго: Эрдэнэс Тавантолгой гэх мэт компаниудад зарах')

doc.add_heading('Санал болгох зам:', level=2)
doc.add_paragraph('Energy Resources-тэй хамтрах:')
doc.add_paragraph('1. Удирдлагатай нээлттэй ярилцах')
doc.add_paragraph('2. "Энэ системийг бусад компанид зарж, компанид орлого оруулж болно" гэж санал болгох')
doc.add_paragraph('3. Revenue sharing тохиролцох (жишээ: 70% танд, 30% компанид)')
doc.add_paragraph('4. Та - техникийн түнш, тэд - бизнесийн түнш')

doc.add_heading('Давуу тал:', level=2)
doc.add_paragraph('- Хууль зүйн асуудалгүй')
doc.add_paragraph('- Energy Resources-ийн нэр хүнд, маркетингийн дэмжлэг')
doc.add_paragraph('- Том компаниуд хоорондоо итгэлцэлтэй худалдаа хийдэг')
doc.add_paragraph('- Эрдэнэс Тавантолгой бол төрийн өмчит - тендер, албан ёсны гэрээ шаарддаг')

doc.add_paragraph()
doc.add_paragraph('─' * 50)
doc.add_paragraph('Энэхүү ярилцлага нь ерөнхий мэдээллийн зорилготой бөгөөд хууль зүйн зөвлөгөө биш болохыг анхаарна уу. Чухал шийдвэр гаргахын өмнө мэргэжлийн хуульчтай зөвлөлдөхийг зөвлөж байна.')

# Save
output_path = 'D:/coal_lims_project/LIMS_Licensing_Discussion.docx'
doc.save(output_path)
print(f'Document saved: {output_path}')
