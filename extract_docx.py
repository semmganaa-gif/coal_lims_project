import sys, io, os

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
    from docx import Document

folder = r'D:\coal_lims_project\burtgel'
for fname in sorted(os.listdir(folder)):
    if fname.endswith('.docx'):
        fpath = os.path.join(folder, fname)
        print()
        print('=' * 80)
        print('FILE: ' + fname)
        print('=' * 80)
        doc = Document(fpath)
        
        print()
        print('--- PARAGRAPHS ---')
        for p in doc.paragraphs:
            if p.text.strip():
                print(p.text)
        
        for i, table in enumerate(doc.tables):
            print()
            print('--- TABLE ' + str(i+1) + ' ---')
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                print(' | '.join(cells))
